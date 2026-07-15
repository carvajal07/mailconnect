from docx import Document
from datetime import datetime
import boto3
import copy
import json
import uuid
import csv
import io

REGION = 'us-east-1'
URL_SQS_EAP = 'https://sqs.us-east-1.amazonaws.com/873837768806/Email_Send-batch-raw-EAP'

#Configurar el cliente de DynamoDB
dynamodb = boto3.resource('dynamodb')

# Configurar el cliente de SQS
sqs = boto3.client('sqs', region_name=REGION)

#Crea el cliente para S3
s3 = boto3.client('s3', region_name=REGION)

global process_detail_id
global formatted_date
global customer_name
global bucket_name
global process_id

table_document = dynamodb.Table('document')

def send_sqs(url_sqs:str,message:list)->None:
    """
    Esta función realiza el envio a las colas de SQS.

    Args:
        url_sqs (str): Url de la cola SQS en AWS
        messages (list): Lista con los mensajes (Maximo 10)

    Returns:
        dict: Nombre de la campaña
    """

    
    try:
        response = sqs.send_message(
            QueueUrl=url_sqs,
            MessageBody=json.dumps(message) if not isinstance(message, str) else message
        )
        print(response)

    except Exception as e:
        # Relanzar: si no se encola el mensaje, SQS debe reintentar / DLQ
        # (antes se tragaba el error y la parte se perdía en silencio).
        print(e)
        raise

    #Validar la posibilidad de reintentos si no se puede encolar

def insert_process_detail(registers:int,part:int,date:str,state:str)->None:
    """
    Función encargada de insertar los detalles de cada parte a la base de datos con su respectivo estado.

    Args:
        registers (int): Cantidad de registros a enviar
        part (int): Indice de la parte
        date (str): Fecha de procesamiento
        state (str): Estado del proceso
        
    Returns:
        None: No retorna resultados
    """   

    table_process_detail = dynamodb.Table(f'{customer_name}_processDetail')

    # Insertar datos en la tabla de detalle de procesos
    table_process_detail.put_item(
        Item={
            'processDetailId': process_detail_id,
            'processId': process_id,
            'registers': registers,
            'part': part,
            'date': date,
            'state': state
        }
    )

def validate_process_detail(part:int)->dict:
    """
    Función encargada de validar el estado de cada parte en la tabla de los detalles.

    Args:
        part (int): Indice de la parte a validar
        
    Returns:
        dict: Informacion de la parte
    """

    table_process_detail = dynamodb.Table(f'{customer_name}_processDetail')
    projection_campaign_expression = 'state, processDetailId'  # Lista de campos a consultar

    response_process_detail = table_process_detail.scan(
        FilterExpression="processId = :value1 and part = :value2",
        ExpressionAttributeValues={":value1": process_id,":value2": part},
        ProjectionExpression=projection_campaign_expression
    )
    return response_process_detail

def download_attachments_data(campaign_id:str)->str:
    projection_document_expression = 'documentPath'  # Lista de campos a consultar

    response_document = table_document.scan(
        FilterExpression="campaignId = :value",
        ExpressionAttributeValues={":value": campaign_id},
        ProjectionExpression=projection_document_expression
    )

    #Revisar porque aca puede ser mas de un adjunto
    items = response_document['Items']
    
    if items:
        for item in items:
            attachment_path = item["documentPath"]

            # Descargar el objeto S3 en un objeto BytesIO
            file_name = attachment_path.split('/')[1]

            #Descargar la plantilla a un directorio temporal
            template_temp_file = f'/tmp/{customer_name}_{formatted_date}_file_name.tmp' 
            s3.download_file(bucket_name, attachment_path, template_temp_file)

            return template_temp_file
    else:
        print("Error, el adjunto para el envio no se encuentra registrado en la tabla de documentos")
        print(f"El id de campaña {campaign_id} no se encontro en la tabla document")

def lambda_handler(event, context):
    """
    Función principal

    Args:
        event (dict): Datos de evento
        context (dict): Datos de contexto
        
    Returns:
        None: Personalizado
    """

    # Procesa TODOS los records del batch SQS (antes solo se leia Records[0],
    # perdiendo el resto si el trigger usa BatchSize>1). Se re-invoca el handler
    # con un record a la vez para reutilizar el flujo existente por-registro.
    _records = event.get("Records") if isinstance(event, dict) else None
    if _records and len(_records) > 1:
        _results = []
        for _rec in _records:
            _results.append(lambda_handler({"Records": [_rec]}, context))
        return _results
    global process_detail_id
    global formatted_date
    global customer_name
    global process_id
    global bucket_name

    # Obtener la fecha y hora actual
    now = datetime.utcnow()
    # Formatear la fecha y hora según un formato específico
    formatted_date = now.strftime("%Y-%m-%dT%H:%M:%S.%f")[:-3] + 'Z'
    id = str(uuid.uuid4())
    process_detail_id = str(uuid.uuid4())
    try:
        # Obtener datos del evento
        body = event["Records"][0]["body"]
        json_body = json.loads(body)
        customer_id = json_body["customerId"]
        customer_name = json_body["customerName"]
        process_id = json_body["processId"]
        campaign_id = json_body["campaignId"]
        attachment = json_body["attachment"]
        from_email = json_body["fromEmail"]
        headers = json_body["headers"]
        template_name = json_body["templateName"]
        part = json_body["part"]
        data = json_body["data"]
        registers = len(data)
        print(f"Customer: {customer_name}")
        print(f"Customer id: {customer_id}")
        print(f"Process id: {process_id}")
        print(f"Campaign id: {campaign_id}")
        print(f"From email: {from_email}")
        print(f"Headers: {headers}")
        print(f"Template name: {template_name}")
        print(f"Parte: {part}")
        print(f"Cantidad registros a procesar: {registers}")

    except Exception as e:
        print(e)
        status = False
        statusCode = 500
        description = "Error no controlado en el servicio"
    else:        
        #Validar el estado de la parte a procesar (Si se encuentra procesando, creando adjuntos o terminada puede ser un error de mensajes duplicados)
        #Debo generar error para evitar realizar envios duplicados
        response_process_detail = validate_process_detail(part)
        if response_process_detail['Items']:
            state = response_process_detail['Items'][0]["state"]
            print(f"La parte {part} del proceso {process_id} ya se encuentra procesando o ha finalizado")
            print(f"El id: {process_detail_id} se encuentra en estado {state}")
            raise ValueError("La parte ya ha sido procesada")
        
        insert_process_detail(registers,part,formatted_date,"Creando adjuntos")
        bucket_name = f'{customer_name}.document'
        template_file = download_attachments_data(campaign_id)
        # Carga la plantilla DOCX original
        docOriginal = Document(template_file)
        headers_list = headers.split(";")
        
        # Itera sobre cada fila del CSV
        for register in data:
            doc = copy.deepcopy(docOriginal) #3.6 segundos con 250 archivos docx
            data_list = register.split(";")
            doc_name = str(data_list[2]) + ".docx"

            # Reemplaza el texto en la plantilla
            for i, value in enumerate(data_list):
                key = headers_list[i]
                # Reemplazar en párrafos
                for p in doc.paragraphs:
                    if key in p.text:
                        p.text = p.text.replace(key, value)
                
                # Reemplazar en tablas
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                if key in p.text:
                                    p.text = p.text.replace(key, value)

                # Reemplazar en encabezados y pies de página
                for section in doc.sections:
                    header = section.header
                    for p in header.paragraphs:
                        if key in p.text:
                            p.text = p.text.replace(key, value)

                    footer = section.footer
                    for p in footer.paragraphs:
                        if key in p.text:
                            p.text = p.text.replace(key, value)

                # Reemplazar en cuadros de texto, notas al pie y notas al final
                for shape in doc.inline_shapes:
                    if shape.text and key in shape.text:
                        shape.text = shape.text.replace(key, value)

            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            s3.upload_fileobj(buffer, Bucket=bucket_name, Key=f'{doc_name}.pdf')
            #s3.put_object(Body=file_content, Bucket='my-bucket', Key='myfile.txt')

        body = {
            "customerId":customer_id,
            "customerName":customer_name,
            "processId":process_id,
            "campaignId":campaign_id,
            "attachment":attachment,
            "fromEmail":from_email,
            "headers":headers,
            "templateName":template_name,
            "part":part,
            "data":data
        }
        send_sqs(URL_SQS_EAP,body)